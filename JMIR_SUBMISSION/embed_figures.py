"""
Automatically embed supplementary figures into MULTIMEDIA_APPENDIX_1.docx
Creates: MULTIMEDIA_APPENDIX_1_COMPLETE.docx (all-in-one file)
"""

from docx import Document
from docx.shared import Inches, Pt
import os

# File paths
INPUT_DOCX = "JMIR_SUBMISSION/MULTIMEDIA_APPENDIX_1_TEMP.docx"
OUTPUT_DOCX = "JMIR_SUBMISSION/MULTIMEDIA_APPENDIX_1_COMPLETE.docx"
FIGURE_FOLDER = "JAMIA_SUBMISSION/Supplementary/"

# Figure mappings - match exact paragraph text with full captions
FIGURES = [
    {
        "name": "FigureS1.png",
        "title": "Supplementary Figure S1: ROC Curves and Threshold Analysis",
        "description": "ROC Curves",
        "caption": "Receiver Operating Characteristic curves for each model showing: ROC curve with Area Under Curve (AUC), Optimal threshold point (Youden's index), Current threshold (0.5) marking, and Sensitivity-specificity trade-off. Results: GPT-4o AUC: 0.502 (95% CI: 0.40-0.60), Gemini-2.0 AUC: 0.515 (95% CI: 0.41-0.62), Qwen-Plus AUC: 0.502 (95% CI: 0.40-0.60). Interpretation: AUC values near 0.5 indicate discrimination ability no better than random chance, consistent with ~50% accuracy findings."
    },
    {
        "name": "FigureS2.png",
        "title": "Supplementary Figure S2: Prediction Probability Distributions",
        "description": "Prediction Distributions",
        "caption": "Histograms showing distribution of averaged prediction probabilities (across 4 runs) for each model, stratified by ground truth (Blue bars: No Disease, Red bars: Disease). Observations: Most predictions cluster near 1.0 (disease present); Minimal separation between distributions for disease vs no-disease cases; Very few predictions below 0.5 threshold; Distribution overlap explains poor discrimination."
    },
    {
        "name": "FigureS3.png",
        "title": "Supplementary Figure S3: Feature Importance Analysis",
        "description": "Feature Correlations",
        "caption": "Correlation heatmap showing relationship between clinical features and model predictions (averaged across all models and runs). Key Correlations: Number of vessels on fluoroscopy (ca): r = 0.42*; ST depression (oldpeak): r = 0.38*; Exercise-induced angina: r = 0.31*; Cholesterol level: r = 0.28*. (* p < 0.05)"
    },
    {
        "name": "FigureS4.png",
        "title": "Supplementary Figure S4: Consistency Score Distributions",
        "description": "Consistency Distributions",
        "caption": "Distribution of consistency scores (proportion of identical predictions across 4 runs) for each model. Results: Mean consistency scores: GPT-4o: 0.985, Gemini-2.0: 0.990, Qwen-Plus: 0.980. Observations: All models show very high consistency (>98%); Most cases have perfect consistency (score = 1.0); Rare inconsistent predictions occur in <5% of cases."
    },
    {
        "name": "FigureS5.png",
        "title": "Supplementary Figure S5: Model Performance Comparison",
        "description": "Model Comparison",
        "caption": "Comprehensive comparison of all models across key performance metrics (Accuracy, Precision, Recall, F1-Score, Consistency). Findings: All models achieve similar accuracy (~49-51%); Recall uniformly high (98-100%); Precision matches accuracy (~49-51%); Consistency exceptionally high (98-99%); No significant differences between models (p > 0.05)."
    },
    {
        "name": "FigureS6.png",
        "title": "Supplementary Figure S6: Threshold Optimization",
        "description": "Threshold Optimization",
        "caption": "Analysis of optimal prediction thresholds for maximizing Youden's index (Sensitivity + Specificity - 1). Results: Optimal thresholds range from 0.52-0.58 across models; Minimal improvement over default 0.5 threshold (Δ accuracy < 2%); Current threshold (0.5) is near-optimal for all models. Interpretation: Threshold adjustment cannot compensate for poor discrimination ability."
    }
]

def embed_figures():
    """Main function to embed all figures"""
    
    print("🔄 Opening document...")
    doc = Document(INPUT_DOCX)
    
    figures_embedded = 0
    paragraphs_to_process = []
    embedded_figures = set()  # Track which figures we've already embedded
    
    # First pass: find all figure placeholder paragraphs (only first occurrence)
    for para_idx, para in enumerate(doc.paragraphs):
        para_text = para.text.strip()
        
        # Check if this paragraph matches a figure title AND is short (placeholder only)
        for fig in FIGURES:
            if para_text == fig["title"] and fig["name"] not in embedded_figures:
                # This is likely the image placeholder paragraph (empty except title)
                # Check if it's a standalone paragraph (not followed by description)
                if para_idx + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[para_idx + 1].text.strip()
                    # If next paragraph starts with "Description:" or "Observations:", 
                    # this is the placeholder we want to replace
                    if next_para.startswith("Description:") or next_para.startswith("Observations:") or para_text == next_para:
                        paragraphs_to_process.append((para_idx, para, fig))
                        embedded_figures.add(fig["name"])  # Mark as processed
                        break
    
    # Second pass: replace placeholders with images
    for para_idx, para, fig in paragraphs_to_process:
        print(f"\n📍 Found placeholder at para {para_idx}: {fig['description']}")
        
        # Build full image path
        image_path = os.path.join(FIGURE_FOLDER, fig["name"])
        
        # Check if image exists
        if not os.path.exists(image_path):
            print(f"   ⚠️  WARNING: Image not found: {image_path}")
            continue
        
        # Clear the paragraph text (remove placeholder)
        para.clear()
        
        # Add the image to the paragraph
        run = para.add_run()
        try:
            run.add_picture(image_path, width=Inches(6.5))
            print(f"   ✅ Embedded: {fig['name']} (6.5 inches wide)")
            
            # Add caption below the image (in the same paragraph)
            para.add_run("\n")
            caption_run = para.add_run(fig["caption"])
            # Format caption: italic, size 10
            caption_run.font.italic = True
            caption_run.font.size = Pt(10)
            
            print(f"   📝 Added caption ({len(fig['caption'])} chars)")
            figures_embedded += 1
        except Exception as e:
            print(f"   ❌ Error embedding {fig['name']}: {str(e)}")
    
    # Save the modified document
    print(f"\n💾 Saving to: {OUTPUT_DOCX}")
    doc.save(OUTPUT_DOCX)
    
    # Get file size
    file_size_mb = os.path.getsize(OUTPUT_DOCX) / (1024 * 1024)
    
    print("\n" + "="*60)
    print("✅ COMPLETE!")
    print("="*60)
    print(f"Figures embedded: {figures_embedded}/6")
    print(f"Output file: {OUTPUT_DOCX}")
    print(f"File size: {file_size_mb:.2f} MB")
    
    if file_size_mb > 10:
        print("\n⚠️  WARNING: File size > 10 MB")
        print("   Consider compressing images in Word:")
        print("   Picture Format → Compress Pictures → 150 ppi")
    else:
        print("\n✅ File size is acceptable for JMIR submission")
    
    print("\n📤 Next step: Upload MULTIMEDIA_APPENDIX_1_COMPLETE.docx to JMIR")
    print("   (Do NOT upload FigureS1-S6.png separately)")

if __name__ == "__main__":
    try:
        embed_figures()
    except FileNotFoundError as e:
        print(f"\n❌ ERROR: File not found")
        print(f"   {str(e)}")
        print("\nPlease check:")
        print(f"  1. {INPUT_DOCX} exists")
        print(f"  2. {FIGURE_FOLDER} contains FigureS1-S6.png")
    except Exception as e:
        print(f"\n❌ ERROR: {str(e)}")
        print("\nIf you encounter issues, use the manual guide:")
        print("  JMIR_SUBMISSION/EMBED_FIGURES_MANUAL_GUIDE.md")
    print("✅ Within JMIR 10 MB limit")
else:
    print("⚠️ WARNING: File exceeds JMIR 10 MB limit - consider compressing images")
